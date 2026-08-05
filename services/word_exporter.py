from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class WordExporter:

    @staticmethod
    def export(report):

        document = Document()

        # ================================
        # Marges
        # ================================

        section = document.sections[0]

        section.top_margin = Cm(2)

        section.bottom_margin = Cm(2)

        section.left_margin = Cm(2)

        section.right_margin = Cm(2)

        # ================================
        # Logo
        # ================================

        try:

            paragraph = document.add_paragraph()

            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            run = paragraph.add_run()

            run.add_picture(
                "static/images/logo_hcp.png",
                width=Cm(2.8)
            )

        except Exception:

            pass


        subtitle = document.add_paragraph()

        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = subtitle.add_run(
            "Haut-Commissariat au Plan\n"
            "Plateforme intelligente d'aide à la décision"
        )

        run.bold = True
        run.font.size = Pt(12)

        # ================================
        # Titre
        # ================================

        title = document.add_heading(
            "Rapport d'aide à la décision",
            level=1
        )

        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # ================================
        # Informations générales
        # ================================

        document.add_heading(
            "Informations générales",
            level=2
        )

        table = document.add_table(
            rows=4,
            cols=2
        )

        table.style = "Table Grid"

        table.cell(0,0).text = "Nom du dataset"
        table.cell(0,1).text = report["filename"]

        table.cell(1,0).text = "Base de données"
        table.cell(1,1).text = report["database"]

        table.cell(2,0).text = "Table PostgreSQL"
        table.cell(2,1).text = report["table"]

        table.cell(3,0).text = "Date"
        table.cell(3,1).text = report["date"]

        # ================================
        # Résumé
        # ================================

        document.add_heading(
            "Résumé exécutif",
            level=2
        )

        document.add_paragraph(
            report["summary"]
        )

        # ================================
        # Vue d'ensemble
        # ================================

        document.add_heading(
            "Vue d'ensemble du jeu de données",
            level=2
        )

        overview = document.add_table(
            rows=4,
            cols=2
        )

        overview.style = "Table Grid"

        overview.cell(0,0).text = "Observations"
        overview.cell(0,1).text = str(report["rows"])

        overview.cell(1,0).text = "Variables"
        overview.cell(1,1).text = str(report["columns"])

        overview.cell(2,0).text = "Variables numériques"
        overview.cell(2,1).text = str(report["numeric_columns"])

        overview.cell(3,0).text = "Variables textuelles"
        overview.cell(3,1).text = str(report["text_columns"])

        # ================================
        # Analyse intelligente
        # ================================

        document.add_heading(
            "Analyse intelligente",
            level=2
        )

        for insight in report["insights"]:

            document.add_paragraph(
                insight,
                style="List Bullet"
            )

        # ================================
        # Recommandations
        # ================================

        document.add_heading(
            "Recommandations",
            level=2
        )

        for recommendation in report["recommendations"]:

            document.add_paragraph(
                recommendation,
                style="List Bullet"
            )

        # ================================
        # Conclusion
        # ================================

        document.add_heading(
            "Conclusion",
            level=2
        )

        document.add_paragraph(
            report["conclusion"]
        )

        # ================================
        # Sauvegarde mémoire
        # ================================
        WordExporter.add_page_number(document)
        buffer = BytesIO()

        document.save(buffer)

        buffer.seek(0)

        return buffer


    @staticmethod
    def add_page_number(document):

        section = document.sections[0]

        footer = section.footer

        paragraph = footer.paragraphs[0]

        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        paragraph.add_run("Page ")

        fld = OxmlElement('w:fldSimple')

        fld.set(qn('w:instr'), 'PAGE')

        paragraph._p.append(fld)